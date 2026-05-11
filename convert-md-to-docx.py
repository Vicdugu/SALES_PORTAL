import os
import re
from docx import Document
from docx.shared import Pt

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    in_list = False

    for line in lines:
        line = line.rstrip('\n')
        
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            continue

        # Headers
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            doc.add_heading(header_match.group(2), level=level)
            continue

        # Lists (simple check for - or *)
        list_match = re.match(r'^[\s]*[\-\*]\s+(.*)', line)
        if list_match:
            doc.add_paragraph(list_match.group(1), style='List Bullet')
            continue

        # Empty lines
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Regular paragraph with basic bold/italic support
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                p.add_run(part[1:-1]).italic = True
            else:
                p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx('FRS_Sales_Till_System.md', 'FRS_Sales_Till_System.docx')
